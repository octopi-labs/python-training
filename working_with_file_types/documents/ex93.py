import os

# Install python-docx
from docx import Document

# open a new document 
current_dir = os.path.dirname(__file__)
filename = os.path.join(current_dir, "test.docx")

# Create a Document object
document = Document()

# Write a paragraph
document.add_paragraph('Hello world!')

paraObj1 = document.add_paragraph('This is a second paragraph.')
paraObj2 = document.add_paragraph('This is a yet another paragraph.')

# insert paragraph before
paraObj1.insert_paragraph_before('Lorem ipsum')

# add text to the end of an existing paragraph
paraObj1.add_run(' This text is being added to the second paragraph.')

# Add heading
document.add_heading('This is Header 0', 0)
document.add_heading('This is Header 1', 1)
document.add_heading('This is Header 2', 2)
document.add_heading('This is Header 3', 3)
document.add_heading('This is Header 4', 4)

# Applying bold and italic
paragraph = document.add_paragraph()
paragraph.add_run('Lorem ipsum ')
paragraph.add_run('dolor').bold = True
paragraph.add_run(' sit amet.')
run = paragraph.add_run('Okay.')
run.bold = True
run.italic = True

# Applying a character style
paragraph = document.add_paragraph('This is Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')

paragraph = document.add_paragraph('This is a Normal text, ')
run = paragraph.add_run('text with emphasis.')
run.style = 'Emphasis'

# Save document
document.save(filename)