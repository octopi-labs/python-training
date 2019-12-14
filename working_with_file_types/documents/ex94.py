import os
from docx import Document

# open a new document 
current_dir = os.path.dirname(__file__)
filename = os.path.join(current_dir, "test.docx")

# Create a Document object
document = Document(filename)

# Add line break
document.add_page_break()

# Add a table
table = document.add_table(rows=2, cols=2)

# Access individual cell
cell = table.cell(0, 1)

# Write to a cell
cell.text = 'Rahul'

# Access row of a table
row = table.rows[1]
row.cells[0].text = 'Foo bar.'
row.cells[1].text = 'And a hearty fizzbuzz to you too sir!'

# Using an iteration
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

# Row count
row_count = len(table.rows)
print("Row Count:", row_count)

# Column count
col_count = len(table.columns)
print("Column Count:", col_count)

# Add row to table
row = table.add_row()
row.cells[0].text = 'new column 1'
row.cells[1].text = 'new column 2'

document.add_paragraph()

# Add data to a new table
# get table data -------------
items = (
    (7, '1024', 'Plush kittens'),
    (3, '2042', 'Furbees'),
    (1, '1288', 'French Poodle Collars, Deluxe'),
)

# add table ------------------
table = document.add_table(1, 3)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description'

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]

# Save document
document.save(filename)