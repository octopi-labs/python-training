import os
from docx import Document

# open a new document 
current_dir = os.path.dirname(__file__)
filename = os.path.join(current_dir, "test.docx")

# Create a Document object
document = Document(filename)

# Total paragraphs
print(len(document.paragraphs))

# first paragraph
print(document.paragraphs[0].text)
# Second paragrpah
print(document.paragraphs[1].text)

# Total number of runs in 10th paragraph
print(len(document.paragraphs[9].runs))

# Each run text
print(document.paragraphs[9].runs[0].text)
print(document.paragraphs[9].runs[1].text)
print(document.paragraphs[9].runs[2].text)
print(document.paragraphs[9].runs[3].text)